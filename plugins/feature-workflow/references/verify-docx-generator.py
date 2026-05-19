#!/usr/bin/env python3
"""
verify-docx-generator.py — Multi-style Report Generator (python-docx fallback)

支援多種報告風格：
  --style intumit    Intumit 品牌（藍+橘，Logo，企業感）
  --style tech-dark  科技深色（深藍底，青綠強調，現代感）

使用方式：
  python3 verify-docx-generator.py \
    --verify .spec/{slug}/verify.md \
    --screenshots .spec/{slug}/screenshots/ \
    --output .spec/{slug}/verify-report.docx \
    --style intumit \
    --logo /path/to/logo.png \
    --cover '{"project":"...","feature":"...","author":"...","date":"...","company":"碩網資訊股份有限公司","version":"v1.0"}'

依賴：python-docx（pip install python-docx）
"""

import argparse, json, os, re, sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("請先安裝 python-docx：python3 -m pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ============================================================
# Theme System
# ============================================================

THEMES = {
    "intumit": {
        "name": "Intumit Brand",
        "primary": "4060AC",
        "primary_rgb": RGBColor(0x40, 0x60, 0xAC),
        "deep": "2B4A90",
        "deep_rgb": RGBColor(0x2B, 0x4A, 0x90),
        "accent": "E38300",
        "accent_rgb": RGBColor(0xE3, 0x83, 0x00),
        "text": "333333",
        "text_rgb": RGBColor(0x33, 0x33, 0x33),
        "muted": "707070",
        "muted_rgb": RGBColor(0x70, 0x70, 0x70),
        "light_bg": "F7F9FC",
        "card_bg": "EAEEF6",
        "table_header_bg": "2B4A90",
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "cover_band": "2B4A90",
        "h1_border_pos": "bottom",
        "h1_border_color": "E38300",
        "body_font": "Roboto",
        "cjk_font": "微軟正黑體",
        "body_size": 10.5,
        "footer_left": "©2026 INTUMIT Inc.  |  Confidential",
    },
    "tech-dark": {
        "name": "Tech Dark",
        "primary": "0F3460",
        "primary_rgb": RGBColor(0x0F, 0x34, 0x60),
        "deep": "16213E",
        "deep_rgb": RGBColor(0x16, 0x21, 0x3E),
        "accent": "00B4D8",
        "accent_rgb": RGBColor(0x00, 0xB4, 0xD8),
        "text": "334155",
        "text_rgb": RGBColor(0x33, 0x41, 0x55),
        "muted": "94A3B8",
        "muted_rgb": RGBColor(0x94, 0xA3, 0xB8),
        "light_bg": "F1F5F9",
        "card_bg": "E2E8F0",
        "table_header_bg": "16213E",
        "table_header_text": RGBColor(0xE2, 0xE8, 0xF0),
        "cover_band": "1A1A2E",
        "h1_border_pos": "left",
        "h1_border_color": "00B4D8",
        "body_font": "Segoe UI",
        "cjk_font": "微軟正黑體",
        "body_size": 10.5,
        "footer_left": "©2026 INTUMIT Inc.  |  Confidential",
    },
    "swiss": {
        "name": "Swiss Minimal",
        "primary": "111111",
        "primary_rgb": RGBColor(0x11, 0x11, 0x11),
        "deep": "000000",
        "deep_rgb": RGBColor(0x00, 0x00, 0x00),
        "accent": "2563EB",
        "accent_rgb": RGBColor(0x25, 0x63, 0xEB),
        "text": "1A1A1A",
        "text_rgb": RGBColor(0x1A, 0x1A, 0x1A),
        "muted": "6B7280",
        "muted_rgb": RGBColor(0x6B, 0x72, 0x80),
        "light_bg": "F9FAFB",
        "card_bg": "F3F4F6",
        "table_header_bg": "111111",
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "cover_band": "000000",
        "h1_border_pos": "none",
        "h1_border_color": "000000",
        "body_font": "Calibri",
        "cjk_font": "微軟正黑體",
        "body_size": 10.5,
        "footer_left": "Confidential",
    },
}

STATUS_MAP = {
    "PASS": ("✅", "通過", RGBColor(0x22, 0x8B, 0x22)),
    "FAIL": ("❌", "未通過", RGBColor(0xCC, 0x00, 0x00)),
    "WARN": ("⚠️", "警告", RGBColor(0xFF, 0x8C, 0x00)),
    "SKIP": ("⏭️", "略過", RGBColor(0x80, 0x80, 0x80)),
    "MANUAL": ("👤", "待人工確認", RGBColor(0x1F, 0x4E, 0x79)),
}

WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ============================================================
# CLI
# ============================================================

def parse_args():
    p = argparse.ArgumentParser(description="verify.md → styled .docx report")
    p.add_argument("--verify", required=True)
    p.add_argument("--screenshots", default=None)
    p.add_argument("--evidence", default=None)
    p.add_argument("--output", required=True)
    p.add_argument("--logo", default=None)
    p.add_argument("--cover", default="{}")
    p.add_argument("--style", default="intumit", choices=list(THEMES.keys()),
                   help="Report style: intumit (brand) or tech-dark (modern)")
    return p.parse_args()

# ============================================================
# verify.md Parser
# ============================================================

def parse_verify_md(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    result = {"summary": {}, "stats": {}, "items": []}

    sm = re.search(r"## 摘要\s*\n(.*?)(?=\n## )", content, re.DOTALL)
    if sm:
        for m in re.finditer(r"\|\s*(.+?)\s*\|\s*(.+?)\s*\|", sm.group(1)):
            k, v = m.group(1).strip(), m.group(2).strip()
            if k not in ("項目", "---"): result["summary"][k] = v

    stm = re.search(r"## 統計\s*\n(.*?)(?=\n## )", content, re.DOTALL)
    if stm:
        for m in re.finditer(r"\|\s*(.+?)\s*\|\s*(\d+)\s*\|", stm.group(1)):
            lb = m.group(1).strip()
            for sk in STATUS_MAP:
                if sk in lb: result["stats"][sk] = int(m.group(2))

    icon_map = {"✅": "PASS", "❌": "FAIL", "⚠️": "WARN", "⏭️": "SKIP", "👤": "MANUAL"}
    for m in re.finditer(r"### \[(\d+)\]\s*(✅|❌|⚠️|⏭️|👤)\s*(.+?)(?=\n### \[|\n---|\Z)", content, re.DOTALL):
        idx, status, body = int(m.group(1)), icon_map.get(m.group(2).strip(), "MANUAL"), m.group(3)
        item = {"index": idx, "status": status, "name": body.split("\n")[0].strip(),
                "type": "", "human_steps": None, "evidence": None, "screenshot": None,
                "fail_reason": None, "skip_reason": None}
        for pat, key in [(r"\*\*類型\*\*：(.+)", "type"), (r"\*\*失敗原因\*\*：(.+)", "fail_reason"),
                         (r"\*\*跳過原因\*\*：(.+)", "skip_reason")]:
            f = re.search(pat, body)
            if f: item[key] = f.group(1).strip()
        shot = re.search(r"\*\*截圖\*\*：(screenshots/.+\.png)", body)
        if shot: item["screenshot"] = shot.group(1).strip()
        hs = re.search(r"<!-- human_steps\s*\n(.*?)-->", body, re.DOTALL)
        if hs:
            steps = {"actions": [], "expected": "", "actual": ""}
            for line in hs.group(1).strip().split("\n"):
                line = line.strip().lstrip("- ")
                if line.startswith("操作："): steps["actions"].append(line[3:].strip())
                elif line.startswith("預期："): steps["expected"] = line[3:].strip()
                elif line.startswith("實際："): steps["actual"] = line[3:].strip()
            item["human_steps"] = steps
        ev = re.search(r"<!-- evidence\s*\n(.*?)-->", body, re.DOTALL)
        if ev: item["evidence"] = ev.group(1).strip()
        result["items"].append(item)
    return result

# ============================================================
# Sensitive Data Masking
# ============================================================

MASK_PATTERNS = [
    (re.compile(r"(Cookie:\s*\w{4})[\w=+/.-]+([\w]{4})"), r"\1****\2"),
    (re.compile(r"(Authorization:\s*\w+\s+\w{4})[\w=+/.-]+"), r"\1****"),
    (re.compile(r"(JSESSIONID=\w{4})[\w=+/.-]+([\w]{4})"), r"\1****\2"),
]

def mask_sensitive(text):
    for pat, repl in MASK_PATTERNS:
        text = pat.sub(repl, text)
    return text

# ============================================================
# Themed Helpers
# ============================================================

def set_shading(cell, hex_color):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))

def set_border(paragraph, position, size, color, space=4):
    pPr = paragraph._element.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(borders)
    borders.append(parse_xml(
        f'<w:{position} {nsdecls("w")} w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'))

def themed_run(p, text, T, size=None, color=None, bold=False, mono=False):
    run = p.add_run(text)
    run.font.size = Pt(size or T["body_size"])
    run.font.color.rgb = color or T["text_rgb"]
    run.font.name = "JetBrains Mono" if mono else T["body_font"]
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
    return run

def themed_table(doc, headers, data, T):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        table._element.insert(0, tblPr)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{T["primary"]}"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{T["primary"]}"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{T["card_bg"]}"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'))
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.color.rgb = T["table_header_text"]
        r.font.size = Pt(10); r.font.name = T["body_font"]
        r._element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
        set_shading(c, T["table_header_bg"])
    for i, rd in enumerate(data):
        row = table.add_row()
        for j, val in enumerate(rd):
            c = row.cells[j]; c.text = ""
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.size = Pt(10)
            r.font.name = T["body_font"]; r._element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
            if i % 2 == 1: set_shading(c, T["light_bg"])
    return table

def add_h1_border(h1_para, T):
    pos = T["h1_border_pos"]
    color = T["h1_border_color"]
    if pos == "none":
        return
    elif pos == "left":
        set_border(h1_para, "left", 24, color, space=8)
    else:
        set_border(h1_para, "bottom", 12, color, space=6)

def add_header_footer(doc, T, logo_path, title_text):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Default header
    header = section.header; header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    if logo_path and os.path.isfile(logo_path):
        hp.add_run().add_picture(logo_path, width=Inches(1.2))
    hp.add_run("\t")
    tr = hp.add_run(title_text)
    tr.font.size = Pt(8); tr.font.color.rgb = T["muted_rgb"]
    tr.font.name = T["body_font"]; tr._element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
    pPr = hp._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9072"/></w:tabs>'))
    set_border(hp, "bottom", 6, T["accent"], space=4)

    # First page header: blank
    fh = section.first_page_header; fh.is_linked_to_previous = False
    for p in fh.paragraphs: p.clear()

    # Default footer
    footer = section.footer; footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    set_border(fp, "top", 4, T["card_bg"], space=4)
    pPr = fp._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9072"/></w:tabs>'))
    r = fp.add_run(T["footer_left"]); r.font.size = Pt(7); r.font.color.rgb = T["muted_rgb"]
    fp.add_run("\t")
    r2 = fp.add_run("Page "); r2.font.size = Pt(7); r2.font.color.rgb = T["muted_rgb"]
    for instr in [" PAGE ", " of ", " NUMPAGES "]:
        if instr.strip() in ("PAGE", "NUMPAGES"):
            fp._element.append(parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>'))
            fp._element.append(parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> {instr.strip()} </w:instrText></w:r>'))
            fp._element.append(parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>'))
        else:
            r3 = fp.add_run(instr); r3.font.size = Pt(7); r3.font.color.rgb = T["muted_rgb"]

    # First page footer: blank
    ff = section.first_page_footer; ff.is_linked_to_previous = False
    for p in ff.paragraphs: p.clear()

# ============================================================
# Cover Page Builders
# ============================================================

def cover_intumit(doc, cover, T, logo_path):
    p = doc.add_paragraph(); set_border(p, "top", 48, T["cover_band"], space=0)
    if logo_path and os.path.isfile(logo_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(36); p.add_run().add_picture(logo_path, width=Inches(2.5))
    p = doc.add_paragraph(); p.space_before = Pt(30)
    set_border(p, "top", 18, T["accent"], space=10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    themed_run(p, cover.get("feature", ""), T, size=22, color=T["deep_rgb"], bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    themed_run(p, "驗 收 報 告", T, size=16, color=T["text_rgb"])
    p = doc.add_paragraph(); set_border(p, "bottom", 18, T["accent"], space=10)
    doc.add_paragraph().space_before = Pt(36)
    themed_table(doc, ["項目", "內容"], [
        ["專案名稱", cover.get("project", "")],
        ["驗證日期", cover.get("date", "")],
        ["版本號", cover.get("version", "")],
        ["承辦單位", cover.get("company", "INTUMIT Inc.（碩網資訊股份有限公司）")],
        ["製作人", cover.get("author", "")],
    ], T)
    p = doc.add_paragraph(); p.space_before = Pt(30)
    set_border(p, "bottom", 48, T["cover_band"], space=0)

def cover_tech_dark(doc, cover, T, logo_path):
    # 深色頂部大色塊
    p = doc.add_paragraph(); set_border(p, "top", 48, T["cover_band"], space=0)
    p = doc.add_paragraph(); set_border(p, "top", 48, T["deep"], space=0)

    if logo_path and os.path.isfile(logo_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(40); p.add_run().add_picture(logo_path, width=Inches(2.0))

    p = doc.add_paragraph(); p.space_before = Pt(40)
    # 青色漸層線
    set_border(p, "top", 18, T["accent"], space=12)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    themed_run(p, cover.get("feature", ""), T, size=24, color=T["primary_rgb"], bold=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    themed_run(p, "ACCEPTANCE TEST REPORT", T, size=11, color=T["muted_rgb"], mono=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    themed_run(p, "驗收報告", T, size=14, color=T["text_rgb"])

    p = doc.add_paragraph()
    set_border(p, "bottom", 18, T["accent"], space=12)

    doc.add_paragraph().space_before = Pt(30)

    # 封面資訊用 code 風格
    info_items = [
        ("PROJECT", cover.get("project", "")),
        ("DATE", cover.get("date", "")),
        ("VERSION", cover.get("version", "")),
        ("COMPANY", cover.get("company", "INTUMIT Inc.（碩網資訊股份有限公司）")),
        ("AUTHOR", cover.get("author", "")),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        themed_run(p, f"  {label}  ", T, size=9, color=T["accent_rgb"], mono=True, bold=True)
        themed_run(p, f"  {value}", T, size=10.5, color=T["text_rgb"])

    p = doc.add_paragraph(); p.space_before = Pt(40)
    set_border(p, "bottom", 48, T["cover_band"], space=0)
    p = doc.add_paragraph(); set_border(p, "bottom", 48, T["deep"], space=0)

def cover_swiss(doc, cover, T, logo_path):
    # 大量留白 + 左對齊大字 + 細黑線分隔 — 國際主義排版
    doc.add_paragraph().space_before = Pt(120)

    # 功能名稱：超大粗體，左對齊
    p = doc.add_paragraph()
    r = p.add_run(cover.get("feature", ""))
    r.font.size = Pt(32); r.font.color.rgb = T["deep_rgb"]; r.bold = True
    r.font.name = T["body_font"]; r._element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])

    # 細黑線
    p = doc.add_paragraph()
    set_border(p, "bottom", 4, T["deep"], space=8)

    # 副標題
    p = doc.add_paragraph(); p.space_before = Pt(8)
    themed_run(p, "驗收報告", T, size=16, color=T["muted_rgb"])

    p = doc.add_paragraph(); p.space_before = Pt(4)
    themed_run(p, "Acceptance Test Report", T, size=11, color=T["muted_rgb"])

    doc.add_paragraph().space_before = Pt(80)

    # 封面資訊：左對齊、label+value 分行、極簡
    info = [
        ("專案", cover.get("project", "")),
        ("日期", cover.get("date", "")),
        ("版本", cover.get("version", "")),
        ("公司", cover.get("company", "INTUMIT Inc.（碩網資訊股份有限公司）")),
        ("製作", cover.get("author", "")),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        themed_run(p, label, T, size=8, color=T["muted_rgb"], bold=True)
        themed_run(p, f"　　{value}", T, size=11, color=T["text_rgb"])

    # 底部細線
    doc.add_paragraph().space_before = Pt(100)
    p = doc.add_paragraph()
    set_border(p, "bottom", 4, T["deep"], space=0)

# ============================================================
# Main Builder
# ============================================================

def build_doc(data, cover, screenshots_dir, evidence_dir, output_path, logo_path, style):
    T = THEMES[style]
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.12)
    section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

    # Default styles
    ns = doc.styles["Normal"]
    ns.font.size = Pt(T["body_size"]); ns.font.name = T["body_font"]
    ns.font.color.rgb = T["text_rgb"]
    ns.element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
    ns.paragraph_format.space_after = Pt(7); ns.paragraph_format.line_spacing = 1.25

    for level, size, color in [(1, 20, T["primary_rgb"]), (2, 14, T["deep_rgb"]), (3, 12, T["text_rgb"])]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = T["body_font"]; hs.font.size = Pt(size); hs.font.color.rgb = color
        hs.font.bold = True; hs.element.rPr.rFonts.set(qn("w:eastAsia"), T["cjk_font"])
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hs.paragraph_format.space_after = Pt(6); hs.paragraph_format.keep_with_next = True

    feature_name = cover.get("feature", "驗收報告")
    add_header_footer(doc, T, logo_path, f"{feature_name} — 驗收報告")

    # Swiss: wider margins, more breathing room
    if style == "swiss":
        section.left_margin = Cm(3.0); section.right_margin = Cm(3.0)
        section.top_margin = Cm(3.0)
        ns.paragraph_format.line_spacing = 1.35

    # ══════ Cover ══════
    if style == "tech-dark":
        cover_tech_dark(doc, cover, T, logo_path)
    elif style == "swiss":
        cover_swiss(doc, cover, T, logo_path)
    else:
        cover_intumit(doc, cover, T, logo_path)
    doc.add_page_break()

    # ══════ 簽核 ══════
    h1 = doc.add_heading("簽核", level=1); add_h1_border(h1, T)
    sign_t = doc.add_table(rows=1, cols=4); sign_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["角色", "姓名", "簽章", "日期"]):
        c = sign_t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True
        r.font.color.rgb = T["table_header_text"]; r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(c, T["table_header_bg"])
    for role, name in [("製作人", cover.get("author", "")), ("審核人", ""), ("客戶確認", "")]:
        row = sign_t.add_row(); row.height = Cm(1.5)
        row.cells[0].text = ""; r = row.cells[0].paragraphs[0].add_run(role)
        r.bold = True; r.font.color.rgb = T["deep_rgb"]; r.font.size = Pt(11)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(row.cells[0], T["light_bg"])
        row.cells[1].text = name; row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = ""; sr = row.cells[2].paragraphs[0].add_run("（簽章）")
        sr.font.color.rgb = T["muted_rgb"]; sr.font.size = Pt(8)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].text = ""; dr = row.cells[3].paragraphs[0].add_run("　　年　　月　　日")
        dr.font.color.rgb = T["muted_rgb"]; dr.font.size = Pt(9)
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Sign borders
    stblPr = sign_t._element.find(qn("w:tblPr"))
    stblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{T["primary"]}"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{T["primary"]}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{T["card_bg"]}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{T["card_bg"]}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{T["card_bg"]}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{T["card_bg"]}"/>'
        f'</w:tblBorders>'))
    doc.add_paragraph()

    # ══════ 測試環境 ══════
    h1 = doc.add_heading("測試環境", level=1); add_h1_border(h1, T)
    themed_table(doc, ["項目", "說明"], [
        ["測試 URL", data["summary"].get("環境", "")],
        ["瀏覽器", "Chromium"],
        ["測試帳號角色", "系統管理員"],
        ["模式", data["summary"].get("模式", "")],
        ["前置條件", "已登入後台管理系統"],
    ], T)

    # ══════ 驗收摘要 ══════
    h1 = doc.add_heading("驗收摘要", level=1); add_h1_border(h1, T)
    stats_data = []
    for key in ["PASS", "FAIL", "WARN", "SKIP", "MANUAL"]:
        if key in data["stats"]:
            icon, label, _ = STATUS_MAP[key]
            stats_data.append([f"{icon} {label}", str(data["stats"][key])])
    themed_table(doc, ["狀態", "數量"], stats_data, T)

    total = sum(data["stats"].values())
    fc = data["stats"].get("FAIL", 0); mc = data["stats"].get("MANUAL", 0); pc = data["stats"].get("PASS", 0)
    if fc == 0 and mc == 0: conc = f"共 {total} 項驗收條件全數通過，建議進入正式上線流程。"
    elif fc > 0: conc = f"共 {total} 項驗收條件，{fc} 項未通過，需修正後重新驗證。"
    else: conc = f"共 {total} 項驗收條件，{pc} 項通過、{mc} 項待人工確認。"
    p = doc.add_paragraph()
    themed_run(p, "結論：", T, size=11, color=T["deep_rgb"], bold=True)
    themed_run(p, conc, T, size=11)

    # ══════ 驗收明細 ══════
    h1 = doc.add_heading("驗收明細", level=1); add_h1_border(h1, T)

    for item in data["items"]:
        icon, label, color = STATUS_MAP.get(item["status"], ("", "", T["text_rgb"]))
        doc.add_heading(f'驗收項目 {item["index"]}：{item["name"]}', level=2)
        p = doc.add_paragraph()
        themed_run(p, "結果：", T, bold=True)
        themed_run(p, f"{label} {icon}", T, bold=True, color=color)

        if item["status"] == "SKIP" and item.get("skip_reason"):
            doc.add_paragraph(f'略過原因：{item["skip_reason"]}'); continue

        if item.get("human_steps"):
            hs = item["human_steps"]
            if hs["actions"]:
                p = doc.add_paragraph(); themed_run(p, "操作步驟：", T, bold=True)
                for i, a in enumerate(hs["actions"], 1): doc.add_paragraph(f"    {i}. {a}")
            if hs["expected"]:
                p = doc.add_paragraph(); themed_run(p, "預期結果：", T, bold=True); themed_run(p, hs["expected"], T)
            if hs["actual"]:
                p = doc.add_paragraph(); themed_run(p, "實際結果：", T, bold=True); themed_run(p, hs["actual"], T)
        elif item.get("fail_reason"):
            p = doc.add_paragraph(); themed_run(p, "失敗原因：", T, bold=True, color=STATUS_MAP["FAIL"][2])
            themed_run(p, item["fail_reason"], T)

        if item.get("evidence"):
            p = doc.add_paragraph(); themed_run(p, "測試紀錄：", T, bold=True)
            p2 = doc.add_paragraph(); r = p2.add_run(mask_sensitive(item["evidence"]))
            r.font.size = Pt(8); r.font.name = "Consolas"

        if item.get("screenshot") and screenshots_dir:
            img = os.path.join(screenshots_dir, os.path.basename(item["screenshot"]))
            if not os.path.isfile(img):
                img = os.path.join(os.path.dirname(screenshots_dir.rstrip("/")), item["screenshot"])
            if os.path.isfile(img):
                doc.add_paragraph()
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img, width=Inches(5.2))
            else:
                doc.add_paragraph(f'（截圖不可用：{item["screenshot"]}）')

    # ══════ 待處理事項 ══════
    ai = [it for it in data["items"] if it["status"] in ("FAIL", "MANUAL")]
    if ai:
        doc.add_page_break()
        h1 = doc.add_heading("待處理事項", level=1); add_h1_border(h1, T)
        themed_table(doc, ["#", "驗收條件", "狀態", "建議"],
            [[str(it["index"]), it["name"], STATUS_MAP[it["status"]][1],
              "修復後重新驗證" if it["status"] == "FAIL" else "手動確認"] for it in ai], T)

    # ══════ 附錄 ══════
    doc.add_page_break()
    h1 = doc.add_heading("附錄", level=1); add_h1_border(h1, T)
    doc.add_heading("版本紀錄", level=2)
    themed_table(doc, ["日期", "版本", "說明"],
        [[cover.get("date", ""), cover.get("version", ""), "初次驗收"]], T)
    doc.add_paragraph()
    doc.add_heading("參考文件", level=2)
    doc.add_paragraph("• 技術規格書：spec.md")
    doc.add_paragraph("• 驗證技術紀錄：verify.md")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"✅ [{T['name']}] 報告已產出：{output_path}")

# ============================================================
# Main
# ============================================================

def main():
    args = parse_args()
    if not os.path.isfile(args.verify):
        print(f"❌ 找不到 verify.md：{args.verify}", file=sys.stderr); sys.exit(1)
    try: cover = json.loads(args.cover)
    except json.JSONDecodeError as e:
        print(f"❌ --cover JSON 解析失敗：{e}", file=sys.stderr); sys.exit(1)
    data = parse_verify_md(args.verify)
    build_doc(data, cover, args.screenshots, args.evidence, args.output, args.logo, args.style)

if __name__ == "__main__":
    main()
